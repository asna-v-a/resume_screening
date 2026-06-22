import re
import os
import spacy
import pdfplumber
import docx

# Lazy load spaCy model with safety fallback
nlp = None
def get_nlp():
    global nlp
    if nlp is None:
        try:
            nlp = spacy.load('en_core_web_sm')
        except OSError:
            import subprocess
            import sys
            # Attempt to download if not found
            subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"], capture_output=True)
            nlp = spacy.load('en_core_web_sm')
    return nlp

def extract_text_from_pdf(filepath):
    """Extracts all text from a PDF file using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {filepath}: {e}")
    return text

def extract_text_from_docx(filepath):
    """Extracts all text from a DOCX file, including paragraphs and tables."""
    text_parts = []
    try:
        doc = docx.Document(filepath)
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
    except Exception as e:
        print(f"Error reading DOCX {filepath}: {e}")
    return "\n".join(text_parts)

def extract_text(filepath):
    """Detects file type and extracts raw text."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(filepath)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(filepath)
    elif ext == '.txt':
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading TXT {filepath}: {e}")
            return ""
    else:
        return ""

def extract_contact_info(text):
    """Extracts email and phone number using regex patterns."""
    # Email regex
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    email = emails[0] if emails else "Not Found"
    
    # Phone number regex (covers various international and local formats)
    phone_pattern = r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'
    phones = re.findall(phone_pattern, text)
    phone = phones[0] if phones else "Not Found"
    
    return email, phone

def extract_name(text):
    """Extracts the candidate name using structural heuristics and spaCy validation."""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    # Check the first 3 non-empty lines - names are always at the very top
    for line in lines[:3]:
        # Skip lines that look like emails, phone numbers, or links
        if '@' in line or '.com' in line or any(c in line for c in '0123456789+()|'):
            continue
        # Skip section headings or metadata
        if line.lower() in ['resume', 'cv', 'curriculum vitae', 'summary', 'contact', 'profile', 'contact info', 'page']:
            continue
        
        # Check if line contains mainly alphabetical characters and is of reasonable length
        clean_line = re.sub(r'[^a-zA-Z\s\.\-\’\u00C0-\u017F]', '', line).strip()
        words = clean_line.split()
        if 1 <= len(words) <= 4 and 2 < len(clean_line) < 35:
            # Simple check: make sure it's not a common tech keyword
            if clean_line.lower() not in ['git', 'docker', 'kubernetes', 'python', 'javascript', 'flask', 'django', 'html', 'css', 'sql']:
                return clean_line
            
    # Fallback to general spaCy PERSON search in the first 300 characters
    spacy_nlp = get_nlp()
    doc = spacy_nlp(text[:300])
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            name = ent.text.strip().replace('\n', ' ')
            if 3 < len(name) < 30 and not any(char in name for char in ['@', '/', '.com', ':', '|']) and not any(c in name for c in '0123456789'):
                if name.lower() not in ['git', 'docker', 'kubernetes', 'python', 'javascript', 'flask', 'django', 'html', 'css', 'sql']:
                    return name
                    
    return "Unknown Candidate"

def extract_skills(text):
    """Cross-references resume text with a detailed industry skills database."""
    skills_database = [
        # Languages
        'python', 'java', 'javascript', 'c++', 'c#', 'c', 'ruby', 'php', 'swift', 'go', 'rust', 'typescript', 'kotlin', 'scala', 'r', 'matlab', 'perl', 'shell', 'bash',
        # Web Frameworks & Libraries
        'react', 'vue', 'angular', 'html', 'css', 'bootstrap', 'tailwind', 'jquery', 'svelte', 'next.js', 'nextjs', 'redux', 'webpack',
        # Backend & Databases
        'flask', 'django', 'fastapi', 'node.js', 'nodejs', 'express', 'spring boot', 'rails', 'laravel', 'asp.net', 'graphql', 'rest api', 'grpc', 'microservices',
        'sql', 'mysql', 'postgresql', 'sqlite', 'mongodb', 'redis', 'cassandra', 'oracle', 'mariadb', 'dynamodb', 'firebase', 'neo4j',
        # Cloud, DevOps & Platforms
        'aws', 'gcp', 'azure', 'docker', 'kubernetes', 'git', 'github', 'gitlab', 'jenkins', 'terraform', 'ansible', 'ci/cd', 'linux', 'nginx', 'apache', 'devops',
        # ML, AI & Data Science
        'machine learning', 'deep learning', 'nlp', 'computer vision', 'tensorflow', 'pytorch', 'keras', 'pandas', 'numpy', 'scikit-learn', 'matplotlib', 'seaborn', 'spark', 'hadoop', 'tableau', 'power bi', 'data science', 'data analysis', 'statistics',
        # Project & Business Methodologies
        'agile', 'scrum', 'project management', 'communication', 'leadership', 'teamwork', 'problem solving', 'time management', 'jira', 'confluence'
    ]
    
    found_skills = []
    text_lower = text.lower()
    
    for skill in skills_database:
        escaped_skill = re.escape(skill)
        # Handle word boundaries dynamically based on symbols (e.g. C++, C#)
        if skill in ['c++', 'c#']:
            pattern = rf'\b{escaped_skill}(?:\b|[^\w]|$)'
        elif skill.startswith('.'):
            pattern = rf'{escaped_skill}\b'
        else:
            pattern = rf'\b{escaped_skill}\b'
            
        if re.search(pattern, text_lower):
            found_skills.append(skill)
            
    # Standardize duplicates (e.g. node.js vs nodejs, next.js vs nextjs)
    synonyms = {
        'nodejs': 'node.js',
        'nextjs': 'next.js'
    }
    cleaned_skills = []
    for skill in found_skills:
        cleaned_skills.append(synonyms.get(skill, skill))
        
    return sorted(list(set(cleaned_skills)))

def parse_sections(text):
    """Categorizes lines of text into sections (Education, Experience, Skills, Projects, Other)."""
    sections = {
        'education': [],
        'experience': [],
        'skills': [],
        'projects': [],
        'other': []
    }
    
    # Common section header keywords
    headings_map = {
        'education': ['education', 'academic background', 'academics', 'qualification', 'qualifications', 'academic profile', 'academic history'],
        'experience': ['experience', 'work experience', 'employment history', 'professional experience', 'work history', 'career history', 'employment', 'work background'],
        'skills': ['skills', 'technical skills', 'core competencies', 'key skills', 'technologies', 'expertise', 'specializations', 'proficiencies'],
        'projects': ['projects', 'academic projects', 'personal projects', 'key projects', 'professional projects', 'selected projects']
    }
    
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    current_section = 'other'
    
    for line in lines:
        lower_line = line.lower().strip(':-#* ')
        
        # Check if the line is a section heading
        matched_section = None
        for sec, keywords in headings_map.items():
            if lower_line in keywords or any(lower_line == kw for kw in keywords):
                matched_section = sec
                break
            # Relaxed match for short lines starting with heading
            if len(lower_line) < 25 and any(lower_line.startswith(kw) for kw in keywords):
                # Ensure it's not a sentence
                if not re.search(r'\b(?:and|or|for|with|in|the)\b', lower_line):
                    matched_section = sec
                    break
        
        if matched_section:
            current_section = matched_section
        else:
            sections[current_section].append(line)
            
    # Format sections as text blocks
    return {k: "\n".join(v) for k, v in sections.items()}

def parse_resume(filepath):
    """High-level orchestrator: reads file, extracts contact, name, skills, sections."""
    raw_text = extract_text(filepath)
    if not raw_text.strip():
        return {
            'name': 'Unknown Candidate',
            'email': 'Not Found',
            'phone': 'Not Found',
            'skills': [],
            'education': '',
            'experience': '',
            'sections': {},
            'raw_text': ''
        }
        
    email, phone = extract_contact_info(raw_text)
    name = extract_name(raw_text)
    skills = extract_skills(raw_text)
    sections = parse_sections(raw_text)
    
    return {
        'name': name,
        'email': email,
        'phone': phone,
        'skills': skills,
        'education': sections.get('education', ''),
        'experience': sections.get('experience', ''),
        'sections': sections,
        'raw_text': raw_text
    }
